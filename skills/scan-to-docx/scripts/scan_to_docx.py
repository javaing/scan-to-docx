#!/usr/bin/env python3
"""
掃描圖片 → .docx
- macOS Vision framework OCR（繁體中文優先）
- 像素分析偵測照片區域（色彩豐富或灰階漸層）
- 圖說行與最近照片配對
- 左欄文字先、右欄文字後
用法:
  每張圖輸出獨立 .docx：
    python3 scan_to_docx.py <圖片或資料夾> <輸出資料夾>
  全部合併成一個 .docx：
    python3 scan_to_docx.py <圖片或資料夾> <輸出路徑.docx>
"""

import sys
import re
from pathlib import Path
import io

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import Quartz
import Vision

CAPTION_PREFIXES = ("•", "●", "‧", "．", "·", "・")


# ── OCR ──────────────────────────────────────────────────────────────────────

def ocr_image(pil_img: Image.Image):
    img_w, img_h = pil_img.size
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG")
    data = buf.getvalue()
    dp = Quartz.CGDataProviderCreateWithData(None, data, len(data), None)
    cg = Quartz.CGImageCreateWithPNGDataProvider(dp, None, False,
                                                  Quartz.kCGRenderingIntentDefault)
    req = Vision.VNRecognizeTextRequest.new()
    req.setRecognitionLanguages_(["zh-Hant", "zh-Hans", "en-US"])
    req.setRecognitionLevel_(Vision.VNRequestTextRecognitionLevelAccurate)
    req.setUsesLanguageCorrection_(True)
    h = Vision.VNImageRequestHandler.alloc().initWithCGImage_options_(cg, {})
    h.performRequests_error_([req], None)

    results = []
    for obs in (req.results() or []):
        top = obs.topCandidates_(1)
        if not top:
            continue
        box = obs.boundingBox()
        x = int(box.origin.x * img_w)
        y = int((1.0 - box.origin.y - box.size.height) * img_h)
        w = int(box.size.width * img_w)
        hpx = int(box.size.height * img_h)
        results.append({
            "text": str(top[0].string()),
            "x": max(0, x), "y": max(0, y), "w": w, "h": hpx,
        })
    return results, img_w, img_h


# ── 像素分析：偵測照片區域 ────────────────────────────────────────────────────

def detect_photo_regions(pil_img, block_size=20, threshold=0.38,
                          min_width=80, min_height=80):
    """
    將頁面分成小方塊，照片方塊的「非白色像素比例」明顯高於文字或空白方塊。
    找出連通的照片方塊群，回傳 list of (x1, y1, x2, y2)。
    """
    gray = pil_img.convert('L')
    W, H = gray.size
    cols = (W + block_size - 1) // block_size
    rows = (H + block_size - 1) // block_size

    # 每個方塊：是否為照片
    grid = []
    for r in range(rows):
        row = []
        for c in range(cols):
            bx1, by1 = c * block_size, r * block_size
            bx2, by2 = min(bx1 + block_size, W), min(by1 + block_size, H)
            block = gray.crop((bx1, by1, bx2, by2))
            hist  = block.histogram()
            total = (bx2 - bx1) * (by2 - by1)
            # 非白色（< 220）佔比
            dark  = sum(hist[:220])
            row.append((dark / total) > threshold if total > 0 else False)
        grid.append(row)

    # BFS 找連通照片方塊群 → 轉成像素邊界框
    visited = [[False] * cols for _ in range(rows)]
    regions = []
    for r0 in range(rows):
        for c0 in range(cols):
            if not grid[r0][c0] or visited[r0][c0]:
                continue
            queue, cells = [(r0, c0)], []
            visited[r0][c0] = True
            while queue:
                r, c = queue.pop(0)
                cells.append((r, c))
                for dr, dc in ((-1,0),(1,0),(0,-1),(0,1)):
                    nr, nc = r+dr, c+dc
                    if (0 <= nr < rows and 0 <= nc < cols
                            and grid[nr][nc] and not visited[nr][nc]):
                        visited[nr][nc] = True
                        queue.append((nr, nc))
            min_r = min(r for r,c in cells); max_r = max(r for r,c in cells)
            min_c = min(c for r,c in cells); max_c = max(c for r,c in cells)
            px1 = min_c * block_size;        py1 = min_r * block_size
            px2 = min((max_c+1)*block_size, W); py2 = min((max_r+1)*block_size, H)
            if (px2-px1) >= min_width and (py2-py1) >= min_height:
                regions.append((px1, py1, px2, py2))

    # 嘗試垂直切割太寬的區域（可能是多張並排的照片）
    regions = _split_wide_regions(gray, sorted(regions, key=lambda r: r[1]),
                                   block_size, threshold)
    return sorted(regions, key=lambda r: r[1])


def _split_wide_regions(gray_img, regions, block_size, threshold,
                         wide_ratio=0.55, scan_step=4):
    """
    對寬度超過頁寬 wide_ratio 的區域，掃描垂直帶找最白的分割線切成兩半。
    """
    W, H = gray_img.size
    result = []
    for rect in regions:
        x1, y1, x2, y2 = rect
        rw = x2 - x1
        if rw < W * wide_ratio:
            result.append(rect)
            continue
        # 在 x 範圍內逐列掃描，找最白（最少暗像素）的 x 位置
        best_x, best_score = x1 + rw // 2, float('inf')
        for x in range(x1 + rw // 5, x2 - rw // 5, scan_step):
            col = gray_img.crop((x, y1, x + scan_step, y2))
            hist = col.histogram()
            total = scan_step * (y2 - y1)
            dark = sum(hist[:220])
            score = dark / total if total > 0 else 0
            if score < best_score:
                best_score, best_x = score, x
        # 只切割如果找到夠白的間隙（暗像素比例 < threshold * 0.6）
        if best_score < threshold * 0.6:
            result.append((x1, y1, best_x, y2))
            result.append((best_x + scan_step, y1, x2, y2))
        else:
            result.append(rect)
    return result


# ── 舊方法：圖說錨點（同欄圖片） ─────────────────────────────────────────────

def build_caption_groups(blocks, img_w):
    """找出 • 開頭的圖說群組及其延續行。"""
    sorted_b = sorted(blocks, key=lambda b: (b["y"], b["x"]))
    groups, i = [], 0
    while i < len(sorted_b):
        b = sorted_b[i]
        if b["text"].startswith(CAPTION_PREFIXES):
            side = "right" if b["x"] > img_w // 2 else "left"
            grp  = {"lines": [b], "side": side,
                    "y_start": b["y"], "y_end": b["y"] + b["h"]}
            j = i + 1
            while j < len(sorted_b):
                nxt = sorted_b[j]
                same_side     = (nxt["x"] > img_w // 2) == (side == "right")
                close_y       = nxt["y"] - grp["y_end"] < 25
                not_new_cap   = not nxt["text"].startswith(CAPTION_PREFIXES)
                if same_side and close_y and not_new_cap:
                    grp["lines"].append(nxt)
                    grp["y_end"] = nxt["y"] + nxt["h"]
                    j += 1
                else:
                    break
            groups.append(grp)
            i = j
        else:
            i += 1
    return groups


def find_image_rect(cap_group, prev_cap_group, img_w, img_h):
    """圖片矩形 = 上一個圖說結束 → 本圖說開始，同側的 x 帶。"""
    side = cap_group["side"]
    x1, x2 = (img_w // 2, img_w) if side == "right" else (0, img_w // 2)
    y1 = (prev_cap_group["y_end"] + 5) if prev_cap_group else 0
    y2 = cap_group["y_start"] - 5
    return (x1, max(0, y1), x2, min(img_h, y2))


# ── 新方法：像素圖說配對（跨欄圖片） ─────────────────────────────────────────

def match_captions_to_regions(regions, blocks, max_gap=130, img_w=None):
    """
    對每個照片區域，收集正下方（距離 < max_gap）的緊鄰文字行作為圖說。
    同欄限制：若 img_w 有傳入，圖說區塊須與圖片同側（左/右欄）。
    回傳 photo entries list 和已被當作圖說的 block id set。
    """
    sorted_b = sorted(blocks, key=lambda b: b["y"])
    used_ids = set()
    entries  = []

    for rect in regions:
        x1, y1, x2, y2 = rect
        # 非跨欄圖片：限制圖說只能來自同側欄
        rect_cx = (x1 + x2) // 2
        if img_w and x1 >= img_w // 2:      # 右欄圖片
            col_min, col_max = img_w // 2, img_w
        elif img_w and x2 <= img_w // 2:    # 左欄圖片
            col_min, col_max = 0, img_w // 2
        else:                               # 跨欄：不限制
            col_min, col_max = 0, img_w or 99999

        cap_lines = []
        cap_col_min, cap_col_max = col_min, col_max  # 找到 • 後可能縮小範圍

        for b in sorted_b:
            bid = id(b)
            if bid in used_ids:
                continue
            if img_w and not (cap_col_min <= b["x"] < cap_col_max):
                continue
            gap = b["y"] - y2
            if gap < -30 or gap > max_gap:   # 允許圖說在圖片底部往上 30px
                continue
            if not cap_lines:
                if b["text"].startswith(CAPTION_PREFIXES):
                    cap_lines.append(b); used_ids.add(bid)
                    # 跨欄圖片：找到 • 後鎖定該行所在欄側
                    if img_w and col_min == 0 and col_max >= img_w:
                        if b["x"] >= img_w // 2:
                            cap_col_min = img_w // 2
                        else:
                            cap_col_max = img_w // 2
            else:
                prev_bot = cap_lines[-1]["y"] + cap_lines[-1]["h"]
                same_level = abs(b["y"] - cap_lines[-1]["y"]) <= 5  # 同一視覺行
                col_w = cap_col_max - cap_col_min
                too_wide = col_w > 0 and b["w"] > col_w * 0.7  # 整欄寬內文，非圖說
                if not too_wide and (same_level or 0 <= b["y"] - prev_bot <= 25):
                    cap_lines.append(b); used_ids.add(bid)
                else:
                    break

        entries.append({
            "rect":     rect,
            "captions": cap_lines,
            "y_insert": y1,
        })

    return entries, used_ids


# ── 寫入 doc ──────────────────────────────────────────────────────────────────

def write_page_to_doc(doc: Document, image_path: str):
    """將單張圖片的 OCR 結果寫入 doc（不存檔）。回傳摘要字串。"""
    pil_img = Image.open(image_path).convert("RGB")
    blocks, img_w, img_h = ocr_image(pil_img)

    if pil_img.size != (img_w, img_h):
        pil_img = pil_img.resize((img_w, img_h), Image.LANCZOS)

    if not blocks:
        buf = io.BytesIO()
        pil_img.save(buf, format="PNG")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return "純圖片"

    # 年代頁偵測：有大字年份（h >= 50）才整頁當圖片，避免誤判圖說內的小字年份
    _year_re = re.compile(r'^(19|20)\d{2}年?$')
    if any(_year_re.match(b["text"].strip()) and b["h"] >= 50 for b in blocks):
        buf = io.BytesIO()
        pil_img.save(buf, format="PNG")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return "年代頁 → 整頁圖片"

    # ── 步驟1：像素偵測所有區域，分跨欄與同欄
    all_rects   = detect_photo_regions(pil_img)
    cross_rects = [r for r in all_rects if r[0] < img_w // 2 and r[2] > img_w // 2]

    def y_overlaps(r, cr, min_ratio=0.5):
        """r 與 cr 的垂直重疊比例是否超過門檻。"""
        overlap = max(0, min(r[3], cr[3]) - max(r[1], cr[1]))
        shorter = min(r[3] - r[1], cr[3] - cr[1])
        return (overlap / shorter) >= min_ratio if shorter > 0 else False

    same_rects  = [r for r in all_rects
                   if r not in cross_rects
                   and not any(y_overlaps(r, cr) for cr in cross_rects)]

    def contained_in(small, big, min_ratio=0.8):
        """small 有 min_ratio 以上面積被 big 包含，視為子區域。"""
        ix1 = max(small[0], big[0]); iy1 = max(small[1], big[1])
        ix2 = min(small[2], big[2]); iy2 = min(small[3], big[3])
        inter = max(0, ix2 - ix1) * max(0, iy2 - iy1)
        area  = (small[2] - small[0]) * (small[3] - small[1])
        return inter / area >= min_ratio if area > 0 else False

    # 同欄：以深色像素比例過濾，並排除被其他區域包含的子區域
    gray_img = pil_img.convert('L')
    all_other = cross_rects + same_rects  # 用來做包含檢查
    same_rects_valid = []
    for r in same_rects:
        x1r, y1r, x2r, y2r = r
        if y2r - y1r >= 50 and x2r - x1r >= 50:
            region = gray_img.crop((x1r, y1r, x2r, y2r))
            hist   = region.histogram()
            total  = (x2r - x1r) * (y2r - y1r)
            dark   = sum(hist[:220])
            if dark / total >= 0.35:
                # 排除被其他更大區域包含的子區域
                if not any(o != r and contained_in(r, o) for o in all_other):
                    same_rects_valid.append(r)

    # ── 步驟2：配對圖說（跨欄與同欄皆用像素偵測結果）
    cross_entries, cross_cap_ids = match_captions_to_regions(cross_rects, blocks,
                                                             img_w=img_w)
    for e in cross_entries:
        e["cross"] = True

    remaining = [b for b in blocks if id(b) not in cross_cap_ids]
    same_entries, same_cap_ids = match_captions_to_regions(same_rects_valid, remaining,
                                                           img_w=img_w)

    # ── 合併，統一排除圖說
    image_entries = cross_entries + same_entries
    all_cap_ids   = cross_cap_ids | same_cap_ids

    # ── 輔助：裁圖並寫入 doc
    def emit_photo(entry):
        x1, y1, x2, y2 = entry["rect"]
        if entry.get("cross"):
            x1, x2 = 0, img_w   # 跨欄圖片左右皆延伸到頁面邊界
            y1 = max(0, y1 - 5)  # 上緣多 5px
        else:
            if entry["captions"]:
                first = entry["captions"][0]
                y2 = max(y2, first["y"] + first["h"])
        y2 = min(y2, img_h)
        if y2 - y1 < 10:
            return
        crop = pil_img.crop((x1, y1, x2, y2))
        buf  = io.BytesIO()
        crop.save(buf, format="PNG")
        buf.seek(0)
        pic_w = Inches(5.0) if entry.get("cross") else Inches(2.5)
        doc.add_picture(buf, width=pic_w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 圖說文字
        for line in entry["captions"]:
            p   = doc.add_paragraph()
            run = p.add_run(line["text"])
            run.font.size      = Pt(10)
            run.font.italic    = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def photo_side(entry):
        x1, _, x2, _ = entry["rect"]
        return "left" if (x1 + x2) // 2 < img_w // 2 else "right"

    # ── 文字方塊（排除所有圖說行）
    left_blocks  = sorted([b for b in blocks
                            if b["x"] < img_w // 2 and id(b) not in all_cap_ids],
                           key=lambda b: b["y"])
    right_blocks = sorted([b for b in blocks
                            if b["x"] >= img_w // 2 and id(b) not in all_cap_ids],
                           key=lambda b: b["y"])

    emitted = set()

    def flush_column(col_blocks, side):
        side_entries = [e for e in image_entries if photo_side(e) == side]
        for b in col_blocks:
            for entry in side_entries:
                eid = id(entry)
                if eid not in emitted and entry["y_insert"] <= b["y"]:
                    emit_photo(entry)
                    emitted.add(eid)
            para = doc.add_paragraph()
            run  = para.add_run(b["text"])
            run.font.size = Pt(12)
        for entry in side_entries:
            eid = id(entry)
            if eid not in emitted:
                emit_photo(entry)
                emitted.add(eid)

    flush_column(left_blocks,  "left")
    flush_column(right_blocks, "right")

    for entry in image_entries:
        if id(entry) not in emitted:
            emit_photo(entry)

    return (f"左欄 {len(left_blocks)} 段 / 右欄 {len(right_blocks)} 段，"
            f"{len(image_entries)} 張圖片（跨欄 {len(cross_entries)}／同欄 {len(same_entries)}）")


def new_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
    return doc


def build_docx(image_path: str, output_dir: str):
    """每張圖輸出一個獨立 .docx。"""
    img_name = Path(image_path).stem
    output_path = Path(output_dir) / f"{img_name}.docx"
    print(f"處理: {Path(image_path).name}", end=" ", flush=True)
    doc = new_doc()
    summary = write_page_to_doc(doc, image_path)
    doc.save(output_path)
    print(f"→ {summary}")


def main():
    if len(sys.argv) < 3:
        print("用法:")
        print("  每張獨立輸出: python3 scan_to_docx.py <圖片或資料夾> <輸出資料夾>")
        print("  合併成一個檔: python3 scan_to_docx.py <圖片或資料夾> <輸出路徑.docx>")
        sys.exit(1)

    src = Path(sys.argv[1])
    out = Path(sys.argv[2])

    exts = {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp"}
    images = [src] if src.is_file() else sorted(
        f for f in src.iterdir() if f.suffix.lower() in exts
    )

    # 判斷模式：輸出路徑是 .docx 檔 → 合併模式
    if out.suffix.lower() == ".docx":
        out.parent.mkdir(parents=True, exist_ok=True)
        print(f"合併模式：共 {len(images)} 張 → {out}\n")
        doc = new_doc()
        for i, img in enumerate(images):
            print(f"處理: {img.name}", end=" ", flush=True)
            try:
                summary = write_page_to_doc(doc, str(img))
                print(f"→ {summary}")
            except Exception as e:
                print(f"→ 錯誤: {e}")
            # 每頁之間加分頁符（最後一頁不加）
            if i < len(images) - 1:
                doc.add_page_break()
        doc.save(out)
    else:
        out.mkdir(parents=True, exist_ok=True)
        print(f"獨立模式：共 {len(images)} 張，輸出到: {out}\n")
        for img in images:
            try:
                build_docx(str(img), str(out))
            except Exception as e:
                print(f"  錯誤: {e}")

    print("\n完成！")


if __name__ == "__main__":
    main()
